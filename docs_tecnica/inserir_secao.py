import copy
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

CAMINHO_ENTRADA = "SIAMP_Documentacao_Tecnica.docx"
CAMINHO_SAIDA = "SIAMP_Documentacao_Tecnica.docx"  # edita in-place

doc = Document(CAMINHO_ENTRADA)

# Localiza a heading "6 SEGURANÇA DA APLICAÇÃO" - novo conteúdo entra
# imediatamente antes dela (final da seção 5, antes da seção 6)
ref = None
for p in doc.paragraphs:
    if p.text.strip() == "6 SEGURANÇA DA APLICAÇÃO" and p.style and p.style.name == "Heading 1":
        ref = p
        break
if ref is None:
    raise RuntimeError("Não encontrou o parágrafo de referência '6 SEGURANÇA DA APLICAÇÃO'")

# Reusa o objeto de estilo direto de um heading já existente no
# documento (5.10), em vez de buscar por nome em doc.styles - evita
# um KeyError inesperado apesar do nome aparecer corretamente listado
# em doc.styles (possível colisão de estilo herdado do template).
estilo_heading2 = None
for p in doc.paragraphs:
    if p.style and p.style.name == "Heading 2":
        estilo_heading2 = p.style
        break
if estilo_heading2 is None:
    raise RuntimeError("Não encontrou nenhum parágrafo com estilo 'Heading 2' para reutilizar.")


def corpo(texto):
    """Parágrafo de corpo de texto, mesmo estilo dos demais (Justify,
    12pt, espaçamento 1.5, 10pt de espaço depois)."""
    p = ref.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(texto)
    r.font.size = Pt(12)
    return p


def heading2(texto):
    """Subtítulo de seção (Heading 2, mesmo estilo de '5.10 ...')."""
    p = ref.insert_paragraph_before()
    p.style = estilo_heading2
    p.text = texto
    return p


def imagem(caminho, largura_cm, altura_cm):
    """Parágrafo centralizado com a imagem, mesmo estilo dos demais."""
    p = ref.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run()
    r.add_picture(caminho, width=Cm(largura_cm), height=Cm(altura_cm))
    return p


def legenda(texto):
    """Legenda de figura, centralizada, itálico, 10pt."""
    p = ref.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(texto)
    r.italic = True
    r.font.size = Pt(10)
    return p


heading2("5.11 Extração de Dados de Ordem de Produção via PDF ou Foto")

corpo(
    "Além do cadastro manual e da importação em lote via CSV/XML "
    "(seção 5.6), o sistema oferece uma terceira via de entrada para "
    "Ordens de Produção: extrair os dados diretamente de um PDF "
    "digitalizado ou de uma foto do documento físico, evitando "
    "redigitação manual de um documento que a empresa já emite pelo "
    "próprio sistema de gestão (MaxManager). A decisão de projeto "
    "mais relevante aqui foi não usar inteligência artificial "
    "generativa para essa extração, apesar de ser a abordagem mais "
    "flexível: como o layout do documento é sempre o mesmo (emitido "
    "pelo mesmo sistema, com posições e rótulos fixos), reconhecimento "
    "de padrão determinístico (expressões regulares sobre o texto "
    "extraído) atende ao caso de uso com igual eficácia, sem o custo "
    "por documento processado, sem exigir configurar uma chave de API "
    "externa, e com comportamento previsível e testável - decisão "
    "confirmada com o usuário antes da implementação."
)

corpo(
    "O maior desafio técnico não é a leitura em si, mas a origem "
    "variável do documento: o mesmo tipo de Ordem de Produção pode "
    "chegar como um PDF gerado digitalmente, com uma camada de texto "
    "real e diretamente extraível, ou como uma digitalização/foto "
    "embutida no PDF, sem texto selecionável algum - caso mais comum "
    "na prática, segundo os documentos reais usados para calibrar o "
    "sistema. O pipeline de extração trata os dois casos com "
    "estratégias diferentes: tenta primeiro ler a camada de texto "
    "nativa do PDF (biblioteca pdfplumber) e, se o resultado tiver "
    "conteúdo substancial, usa esse texto diretamente - muito mais "
    "confiável, sem nenhuma das confusões de caractere que OCR "
    "introduz. Só quando essa tentativa falha (PDF sem texto "
    "selecionável) ou quando o arquivo enviado é uma foto (que nunca "
    "tem texto embutido) o sistema recorre a OCR (Tesseract, com "
    "modelo de idioma português), convertendo a página em imagem e "
    "lendo o texto a partir dela."
)

imagem("18-ordens-producao-importar-pdf.png", 15, 10.38)
legenda("Figura 18 – Card de importação de Ordem de Produção via PDF ou foto. Fonte: elaborado pelo autor.")

corpo(
    "O sistema nunca cadastra a Ordem de Produção diretamente a "
    "partir do documento extraído - o resultado sempre pré-preenche o "
    "mesmo formulário manual de cadastro (incluindo os campos de "
    "seleção de peça e máquina, marcados automaticamente quando o "
    "código extraído bate com um cadastro existente), com um aviso "
    "visível pedindo revisão antes de salvar. Essa exigência de "
    "confirmação humana é deliberada: mesmo com o layout fixo, "
    "extração automática de documento escaneado ou fotografado nunca "
    "é totalmente confiável, e o custo de um erro não percebido "
    "(quantidade errada, peça trocada) é alto o suficiente para "
    "justificar esse passo extra, mesmo à custa de alguma "
    "conveniência."
)

imagem("19-ordens-producao-extraido.png", 15, 9.58)
legenda("Figura 19 – Formulário pré-preenchido a partir de um PDF real, pronto para revisão. Fonte: elaborado pelo autor.")

corpo(
    "A calibração do sistema de reconhecimento foi feita "
    "iterativamente contra dois documentos reais fornecidos pela "
    "empresa, com características propositalmente diferentes - um PDF "
    "só-imagem (sem texto selecionável) e um PDF com texto nativo -, "
    "revelando problemas que não apareceriam testando com apenas um "
    "exemplo ou com dados sintéticos. O número da Ordem de Produção "
    "(identificador único, o campo mais crítico) chegou a sair "
    "completamente corrompido no OCR de página inteira: o Tesseract "
    "confundiu o bloco de fonte grande do número com o texto pequeno "
    "do cabeçalho da empresa ao lado, misturando os dois. A solução "
    "foi recortar computacionalmente só a região da imagem onde esse "
    "número aparece (coordenadas relativas à página, não em pixels "
    "absolutos, para funcionar em qualquer resolução de digitalização) "
    "e reler apenas esse recorte com uma configuração de OCR para "
    "linha única de texto - abordagem que passou a funcionar de forma "
    "perfeitamente confiável nos testes."
)

corpo(
    "Outro problema recorrente, tanto no caminho de OCR quanto no de "
    "texto nativo do PDF, foi o vazamento de campos vizinhos: o "
    "documento tem um layout em duas colunas, e tanto a extração de "
    "texto do PDF quanto o OCR tendem a linearizar esse layout em uma "
    "única sequência de texto, fazendo o valor de um campo (ex.: "
    "'Tipo da OP') absorver o rótulo do campo vizinho posicionado na "
    "mesma altura da página (ex.: 'LOTE'). Corrigido com limites de "
    "parada específicos por campo, identificados um a um ao comparar "
    "o resultado da extração com o conteúdo real do documento. Um "
    "terceiro problema, mais sutil, surgiu de uma anotação escrita à "
    "mão sobre um dos documentos de teste: o OCR leu fragmentos dessa "
    "anotação como se fossem parte do texto impresso vizinho, exigindo "
    "uma heurística de limpeza que reconhece sequências de tokens "
    "curtos e desconexos como ruído - tomando o cuidado de não cortar "
    "conteúdo legítimo igualmente curto (uma sigla de duas letras "
    "seguida de hífen, por exemplo), erro que só foi percebido ao "
    "testar contra o segundo documento real, com um valor desse tipo "
    "genuíno no meio do texto."
)

corpo(
    "O episódio reforça um princípio já observado em outras partes "
    "deste projeto (seção 9): validar contra dados reais, e mais de "
    "uma amostra real quando possível, encontra classes inteiras de "
    "problema que testes com dados idealizados não revelam. A suíte "
    "de testes automatizados reflete essa mesma disciplina - além de "
    "testes unitários sobre a lógica de reconhecimento (usando o "
    "texto já extraído dos documentos reais como referência fixa, "
    "para não depender de rodar OCR de verdade a cada execução da "
    "suíte), há testes de integração que exercitam o pipeline "
    "completo, incluindo OCR de verdade, contra os dois documentos "
    "originais - garantindo que uma mudança futura no código não "
    "regrida silenciosamente um caso que já funcionava."
)

doc.save(CAMINHO_SAIDA)
print("Seção 5.11 inserida com sucesso.")
